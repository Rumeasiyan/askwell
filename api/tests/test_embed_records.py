"""Embedding, end to end, against a real Postgres. `M1-INDEX-ING-032`.

Mirrors `test_chunk_records.py`'s own pattern: documents go through the real
`add()`, `extract.run` and `chunk.run` path, so what is under test is the
installed `embed` stage itself against real chunks — not a stand-in. The
inference process is a real Unix socket serving canned vectors, matching
`test_inference_client.py`'s own reasoning: mocking httpx tests the mock.
"""

import asyncio
import json
import re
import uuid
from collections.abc import AsyncIterator, Iterator
from pathlib import Path
from typing import Any

import docx
import pytest
import pytest_asyncio
from sqlalchemy import text
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine

from askwell import chunk as chunk_module
from askwell import embed as embed_module
from askwell import extract, ingest
from askwell.config import Settings
from askwell.db.engine import session_scope
from askwell.ingest import Work

from .test_ingest_records import TABLES, nominate, recorded

pytestmark = pytest.mark.requires_db

# The width the test database's `chunks.embedding` column was actually
# created at — the migration reads `ASKWELL_EMBEDDING_DIMENSIONS` at the time
# it runs (`conftest_db.py`), and nothing in this suite sets it, so it is
# `Settings`'s own default. A vector of any other length is what
# `test_a_configured_width_that_does_not_match_the_column_is_refused` sends.
DEPLOYED_DIMENSIONS = Settings.model_fields["embedding_dimensions"].default


@pytest.fixture
def async_url(database_url: str) -> str:
    return database_url.replace("postgresql://", "postgresql+psycopg://", 1)


@pytest_asyncio.fixture
async def factory(async_url: str) -> AsyncIterator[async_sessionmaker[AsyncSession]]:
    engine = create_async_engine(async_url)
    sessions = async_sessionmaker(engine, expire_on_commit=False)
    async with sessions() as opened:
        await opened.execute(text(f"TRUNCATE {TABLES} CASCADE"))
        await opened.commit()
    yield sessions
    async with sessions() as opened:
        await opened.execute(text(f"TRUNCATE {TABLES} CASCADE"))
        await opened.commit()
    await engine.dispose()


@pytest_asyncio.fixture
async def session(
    factory: async_sessionmaker[AsyncSession],
) -> AsyncIterator[AsyncSession]:
    async with factory() as opened:
        yield opened
        await opened.rollback()


@pytest.fixture(autouse=True)
def _fast_embed_retries(monkeypatch: pytest.MonkeyPatch) -> None:
    """These tests exercise real retry-until-exhausted paths; a real backoff
    would make them correct and slow. `test_embed.py` proves the backoff
    itself is linear and grows with the attempt."""
    monkeypatch.setattr(embed_module, "EMBED_BATCH_RETRY_DELAY_SECONDS", 0.0)


@pytest.fixture
def unreachable_queue(settings: Settings, monkeypatch: pytest.MonkeyPatch) -> Iterator[Settings]:
    sent: list[uuid.UUID] = []

    async def fake_dispatch(
        _settings: Settings,
        document_ids: list[uuid.UUID],
        **_kwargs: object,
    ) -> int:
        sent.extend(document_ids)
        return len(document_ids)

    monkeypatch.setattr(ingest, "dispatch", fake_dispatch)
    yield settings


def _ready(directory: Path) -> None:
    directory.mkdir(parents=True, exist_ok=True)
    (directory / "state.json").write_text(
        json.dumps({"state": "ready", "model": "a-model.gguf", "acceleration": "cpu"}),
        encoding="utf-8",
    )


class _EmbeddingStub:
    """Answers `/v1/embeddings` with one fixed-width vector per input text.

    Records the size of every batch it was sent, which is how
    `test_the_batch_size_is_bounded_by_configuration` proves a large document
    was actually split rather than sent in one call.
    """

    def __init__(self, dimensions: int) -> None:
        self.dimensions = dimensions
        self.batch_sizes: list[int] = []

    async def handle(self, reader: asyncio.StreamReader, writer: asyncio.StreamWriter) -> None:
        head = await reader.readuntil(b"\r\n\r\n")
        length = 0
        match = re.search(rb"content-length:\s*(\d+)", head, re.I)
        if match:
            length = int(match.group(1))
        payload = json.loads(await reader.readexactly(length) if length else b"{}")
        texts = payload.get("input", [])
        self.batch_sizes.append(len(texts))

        body = json.dumps(
            {"data": [{"embedding": [0.01] * self.dimensions} for _text in texts]}
        ).encode()
        writer.write(
            f"HTTP/1.1 200 X\r\nContent-Type: application/json\r\n"
            f"Content-Length: {len(body)}\r\nConnection: close\r\n\r\n".encode()
            + body
        )
        await writer.drain()
        writer.close()


@pytest_asyncio.fixture
async def inference(
    unreachable_queue: Settings, tmp_path: Path
) -> AsyncIterator[tuple[Settings, _EmbeddingStub]]:
    """A running inference socket the settings under test actually points at."""
    socket_path = tmp_path / "inference.sock"
    _ready(tmp_path)
    stub = _EmbeddingStub(DEPLOYED_DIMENSIONS)
    server = await asyncio.start_unix_server(stub.handle, path=str(socket_path))
    configured = unreachable_queue.model_copy(update={"inference_socket": socket_path})
    yield configured, stub
    server.close()
    await server.wait_closed()


def _write_docx_with_a_rate_table(path: Path) -> None:
    document = docx.Document()
    document.add_heading("Renewal Terms", level=1)
    document.add_paragraph("Either party may terminate on ninety days written notice.")
    table = document.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "Tier"
    table.rows[0].cells[1].text = "Monthly rate"
    table.rows[1].cells[0].text = "Standard"
    table.rows[1].cells[1].text = "199.00"
    table.rows[2].cells[0].text = "Premium"
    table.rows[2].cells[1].text = "349.00"
    document.save(str(path))


async def _chunked_document(
    factory: async_sessionmaker[AsyncSession],
    session: AsyncSession,
    tmp_path: Path,
    filename: str,
    settings: Settings,
    body: bytes | None = None,
) -> tuple[uuid.UUID, Work]:
    """A real document, extracted and chunked for real — everything `embed`
    is handed once `chunk` has already run, built directly rather than
    through `ingest.process` so these tests can call `embed.run` on their own
    and inspect exactly what it wrote."""
    if body is not None:
        (tmp_path / filename).write_bytes(body)
    else:
        _write_docx_with_a_rate_table(tmp_path / filename)

    await nominate(session, str(tmp_path))
    documents = await recorded(session, tmp_path, filename)
    document_id = documents[0]

    async def report(_done: int, _total: int) -> None:
        return None

    mime = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if filename.endswith(".docx")
        else "text/plain"
    )
    work = Work(
        document_id=document_id,
        source_id=uuid.uuid4(),
        path=str(tmp_path / filename),
        filename=filename,
        mime=mime,
        sha256="0" * 64,
    )
    await extract.run(work, report, factory, settings)
    await chunk_module.run(work, report, factory, settings)
    return document_id, work


async def _embedding_rows(session: AsyncSession, document_id: uuid.UUID) -> list[tuple[Any, Any]]:
    rows = (
        await session.execute(
            text(
                "SELECT ordinal, embedding IS NULL, vector_dims(embedding) FROM chunks "
                "WHERE document_id = :id ORDER BY ordinal"
            ),
            {"id": document_id},
        )
    ).all()
    return [tuple(row) for row in rows]


# --- the acceptance criteria --------------------------------------------------


async def test_chunks_receive_embeddings_of_the_configured_dimension(
    factory: async_sessionmaker[AsyncSession],
    session: AsyncSession,
    tmp_path: Path,
    inference: tuple[Settings, _EmbeddingStub],
) -> None:
    configured, _stub = inference
    document_id, work = await _chunked_document(
        factory, session, tmp_path, "rates.docx", configured
    )

    async def report(_done: int, _total: int) -> None:
        return None

    await embed_module.run(work, report, factory, configured)

    async with session_scope(factory) as scoped:
        rows = await _embedding_rows(scoped, document_id)
    assert rows  # the table has chunks worth checking at all
    for _ordinal, is_null, dimensions in rows:
        assert is_null is False
        assert dimensions == DEPLOYED_DIMENSIONS


async def test_a_document_reaches_ready_only_once_every_chunk_is_embedded(
    factory: async_sessionmaker[AsyncSession],
    session: AsyncSession,
    tmp_path: Path,
    inference: tuple[Settings, _EmbeddingStub],
) -> None:
    configured, _stub = inference
    await nominate(session, str(tmp_path))
    _write_docx_with_a_rate_table(tmp_path / "rates.docx")
    documents = await recorded(session, tmp_path, "rates.docx")

    outcome = await ingest.process(factory, configured, documents[0])
    assert outcome == "done"

    status = (
        await session.execute(
            text("SELECT status FROM documents WHERE id = :id"), {"id": documents[0]}
        )
    ).scalar_one()
    assert status == "ready"

    rows = await _embedding_rows(session, documents[0])
    assert rows
    assert all(is_null is False for _ordinal, is_null, _dimensions in rows)


async def test_the_batch_size_is_bounded_by_configuration(
    factory: async_sessionmaker[AsyncSession],
    session: AsyncSession,
    tmp_path: Path,
    inference: tuple[Settings, _EmbeddingStub],
) -> None:
    configured, stub = inference
    small_batches = configured.model_copy(update={"embedding_batch_size": 1})
    sentence = "Either party may terminate this agreement on ninety days written notice. "
    document_id, work = await _chunked_document(
        factory, session, tmp_path, "note.txt", configured, body=(sentence * 400).encode()
    )

    chunk_count = (
        await session.execute(
            text("SELECT count(*) FROM chunks WHERE document_id = :id"), {"id": document_id}
        )
    ).scalar_one()
    assert chunk_count > 1  # otherwise this test proves nothing about batching

    async def report(_done: int, _total: int) -> None:
        return None

    await embed_module.run(work, report, factory, small_batches)

    assert stub.batch_sizes == [1] * chunk_count


async def test_an_inference_process_that_never_answers_leaves_the_document_unindexed(
    factory: async_sessionmaker[AsyncSession],
    session: AsyncSession,
    tmp_path: Path,
    unreachable_queue: Settings,
) -> None:
    """The ticket's own edge case: nothing is embedded, nothing is lost —
    the document simply is not `ready`, and the failure names the stage."""
    await nominate(session, str(tmp_path))
    _write_docx_with_a_rate_table(tmp_path / "rates.docx")
    documents = await recorded(session, tmp_path, "rates.docx")

    # `unreachable_queue`'s socket points nowhere — no supervisor has ever
    # reported in, matching a machine that has not started inference yet.
    outcome = await ingest.process(factory, unreachable_queue, documents[0])
    assert outcome == "failed"

    row = (
        await session.execute(
            text(
                "SELECT j.state, j.stage, j.error, d.status FROM ingest_jobs j "
                "JOIN documents d ON d.id = j.document_id WHERE j.document_id = :id"
            ),
            {"id": documents[0]},
        )
    ).one()
    assert row[1] == "embed"
    assert "InferenceUnavailable" in row[2]
    assert row[3] != "ready"

    rows = await _embedding_rows(session, documents[0])
    assert rows
    assert all(is_null is True for _ordinal, is_null, _dimensions in rows)


async def test_a_failure_exhausted_after_retries_is_visible_and_the_retry_works(
    factory: async_sessionmaker[AsyncSession],
    session: AsyncSession,
    tmp_path: Path,
    inference: tuple[Settings, _EmbeddingStub],
    unreachable_queue: Settings,
) -> None:
    """`MAX_ATTEMPTS` failures without inference, then a retry once it is up —
    the same shape `test_ingest_records.py` already proves for extraction,
    exercised here for the stage this ticket owns."""
    await nominate(session, str(tmp_path))
    _write_docx_with_a_rate_table(tmp_path / "rates.docx")
    documents = await recorded(session, tmp_path, "rates.docx")

    outcome = None
    for _ in range(ingest.MAX_ATTEMPTS):
        outcome = await ingest.process(factory, unreachable_queue, documents[0])
    assert outcome == "failed"

    row = (
        await session.execute(
            text(
                "SELECT j.state, j.error, d.status FROM ingest_jobs j "
                "JOIN documents d ON d.id = j.document_id WHERE j.document_id = :id"
            ),
            {"id": documents[0]},
        )
    ).one()
    assert row[0] == "failed"
    assert row[1] is not None
    assert row[2] == "attention"

    configured, _stub = inference
    retried = await ingest.retry(session, documents[0], configured)
    await session.commit()
    assert retried.retried

    outcome = await ingest.process(factory, configured, documents[0])
    assert outcome == "done"

    status = (
        await session.execute(
            text("SELECT status FROM documents WHERE id = :id"), {"id": documents[0]}
        )
    ).scalar_one()
    assert status == "ready"


async def test_an_empty_chunk_is_rejected_rather_than_embedded(
    factory: async_sessionmaker[AsyncSession],
    session: AsyncSession,
    tmp_path: Path,
    inference: tuple[Settings, _EmbeddingStub],
) -> None:
    """The chunker guarantees this cannot happen; this is the second line the
    ticket's own edge case asks for, proved by forcing the case the chunker
    itself would never produce."""
    configured, _stub = inference
    document_id, work = await _chunked_document(
        factory, session, tmp_path, "rates.docx", configured
    )

    async with session_scope(factory) as scoped:
        await scoped.execute(
            text(
                "INSERT INTO chunks (id, document_id, ordinal, content) "
                "VALUES (:id, :document_id, -1, '   ')"
            ),
            {"id": uuid.uuid4(), "document_id": document_id},
        )

    async def report(_done: int, _total: int) -> None:
        return None

    with pytest.raises(embed_module.EmptyChunk):
        await embed_module.run(work, report, factory, configured)

    async with session_scope(factory) as scoped:
        rows = await _embedding_rows(scoped, document_id)
    # Nothing was embedded — an all-or-nothing batch would have written some
    # chunks and left the empty one out, which is exactly the half-indexed
    # state this ticket exists to prevent.
    assert all(is_null is True for _ordinal, is_null, _dimensions in rows)


# --- the dimension check ------------------------------------------------------


async def test_a_matching_configured_dimension_passes(
    session: AsyncSession, settings: Settings
) -> None:
    await embed_module.check_dimension(session, settings)


async def test_a_configured_width_that_does_not_match_the_column_is_refused(
    session: AsyncSession, settings: Settings
) -> None:
    mismatched = settings.model_copy(update={"embedding_dimensions": DEPLOYED_DIMENSIONS + 1})
    with pytest.raises(embed_module.EmbeddingDimensionMismatch, match=r"but chunks\.embedding"):
        await embed_module.check_dimension(session, mismatched)


async def test_worker_startup_refuses_to_start_on_a_dimension_mismatch(
    factory: async_sessionmaker[AsyncSession],
    settings: Settings,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """`docs/build-plan.md`'s own edge case: refused at startup, not per
    batch. `askwell.worker.startup` is the caller; this proves it actually
    lets the error through rather than swallowing it the way an ordinary
    "Postgres is not up yet" is swallowed."""
    from askwell import worker

    mismatched = settings.model_copy(update={"embedding_dimensions": DEPLOYED_DIMENSIONS + 1})
    monkeypatch.setattr(worker, "build_engine", lambda _settings: None)
    monkeypatch.setattr(worker, "session_factory", lambda _engine: factory)

    with pytest.raises(embed_module.EmbeddingDimensionMismatch):
        await worker.startup({"settings": mismatched})
