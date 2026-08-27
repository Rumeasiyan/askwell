"""Word extraction's own rules, without a database.

`M1-EXTRACT-ING-027`. What is pure here — heading and list formatting, table
markers, page-break and revision detection — needs only a `python-docx`
object in memory. The database write is covered end to end in
`test_extract_office_records.py`, alongside `test_extract_pdf_records.py`'s
own pattern for `extract_pdf`.
"""

import docx
from docx.oxml.ns import qn

from askwell.extract_docx import (
    _has_page_break,
    _has_revisions,
    _iter_block_items,
    _paragraph_text,
    _table_text,
)


def test_a_heading_style_becomes_a_markdown_style_prefix() -> None:
    document = docx.Document()
    document.add_heading("Renewal Terms", level=2)
    paragraph = document.paragraphs[0]
    assert _paragraph_text(paragraph) == "## Renewal Terms"


def test_an_ordinary_paragraph_is_unchanged() -> None:
    document = docx.Document()
    document.add_paragraph("Either party may terminate on ninety days written notice.")
    paragraph = document.paragraphs[0]
    assert _paragraph_text(paragraph) == "Either party may terminate on ninety days written notice."


def test_a_bulleted_list_item_is_marked() -> None:
    document = docx.Document()
    document.add_paragraph("First item", style="List Bullet")
    paragraph = document.paragraphs[0]
    assert _paragraph_text(paragraph) == "- First item"


def test_a_blank_paragraph_contributes_nothing() -> None:
    document = docx.Document()
    document.add_paragraph("   ")
    paragraph = document.paragraphs[0]
    assert _paragraph_text(paragraph) == ""


def test_a_table_becomes_bracketed_rows() -> None:
    document = docx.Document()
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "Price"
    table.rows[1].cells[0].text = "Widget"
    table.rows[1].cells[1].text = "9.99"
    body = _table_text(table)
    assert body.splitlines() == ["[TABLE]", "Item | Price", "Widget | 9.99", "[/TABLE]"]


def test_a_document_with_no_explicit_page_break_has_none_detected() -> None:
    document = docx.Document()
    document.add_paragraph("Only paragraph.")
    assert _has_page_break(document.paragraphs[0]) is False


def test_an_explicit_page_break_is_detected_on_its_own_paragraph() -> None:
    document = docx.Document()
    document.add_page_break()
    assert _has_page_break(document.paragraphs[0]) is True


def test_paragraphs_and_tables_are_walked_in_document_order() -> None:
    document = docx.Document()
    document.add_paragraph("Before the table.")
    document.add_table(rows=1, cols=1)
    document.add_paragraph("After the table.")
    kinds = [type(block).__name__ for block in _iter_block_items(document)]
    assert kinds == ["Paragraph", "Table", "Paragraph"]


def test_a_document_with_no_tracked_changes_reports_none() -> None:
    document = docx.Document()
    document.add_paragraph("Nothing has been edited.")
    assert _has_revisions(document) is False


def test_a_tracked_insertion_is_detected() -> None:
    """`python-docx` has no API for tracked changes, so the insertion is
    built by hand: a `<w:ins>` wrapping a run, exactly the shape Word writes
    when "Track Changes" is on and someone adds a sentence."""
    document = docx.Document()
    paragraph = document.add_paragraph("Some text.")
    ins = paragraph._p.makeelement(qn("w:ins"), {})
    run = paragraph.add_run(" Added later.")
    run._element.getparent().remove(run._element)
    ins.append(run._element)
    paragraph._p.append(ins)
    assert _has_revisions(document) is True
