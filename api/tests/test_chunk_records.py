"""Structure-aware chunking, end to end, against a real Postgres.
`M1-INDEX-ING-031`.

Mirrors `test_extract_office_records.py`'s pattern: files go through the
real `add()` path and then the real `ingest.process`, so what is under test
is the installed `chunk` stage itself — a real `document_pages` row in, real
`chunks` rows out — not a stand-in.
"""

import uuid
from collections.abc import AsyncIterator, Iterator
from pathlib import Path

import docx
import pytest
import pytest_asyncio
from sqlalchemy import text
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine

from askwell import chunk as chunk_module
from askwell import extract, ingest
from askwell.config import Settings
from askwell.db.engine import session_scope
from askwell.ingest import Stage, Work

from .test_ingest_records import TABLES, nominate, recorded

pytestmark = pytest.mark.requires_db

# Duplicated rather than imported, matching `test_extract_office_records.py`'s
# own note: a fixture reused across modules by import is flagged by ruff
# (F811) the moment a test's own parameter shadows the imported name.


@pytest.fixture
def async_url(database_url: str) -> str:
    return database_url.replace("postgresql://", "postgresql+psycopg://", 1)


@pytest_asyncio.fixture
async def factory(async_url: str) -> AsyncIterator[async_sessionmaker[AsyncSession]]:
    engine = create_async_engine(async_url)
    sessions = async_sessionmaker(engine, expire_on_commit=False)
    async with sessions() as opened:
        await opened.execute(text(f"TRUNCATE {TABLES} CASCADE"))
        await opened.commit()
    yield sessions
    async with sessions() as opened:
        await opened.execute(text(f"TRUNCATE {TABLES} CASCADE"))
        await opened.commit()
    await engine.dispose()


@pytest_asyncio.fixture
async def session(
    factory: async_sessionmaker[AsyncSession],
) -> AsyncIterator[AsyncSession]:
    async with factory() as opened:
        yield opened
        await opened.rollback()


@pytest.fixture
def unreachable_queue(settings: Settings, monkeypatch: pytest.MonkeyPatch) -> Iterator[Settings]:
    sent: list[uuid.UUID] = []

    async def fake_dispatch(
        _settings: Settings,
        document_ids: list[uuid.UUID],
        **_kwargs: object,
    ) -> int:
        sent.extend(document_ids)
        return len(document_ids)

    monkeypatch.setattr(ingest, "dispatch", fake_dispatch)
    # This module is about `chunk`, not `embed` — `M1-INDEX-ING-032` made
    # `embed` real and it needs a running inference process none of these
    # tests stand one up for. Frozen at real `extract` + `chunk` with `embed`
    # left unbuilt, matching every test's own assertion that a document
    # "parks" once chunking is done. `test_embed_records.py` covers `embed`.
    monkeypatch.setattr(
        ingest,
        "STAGES",
        (
            Stage("extract", "M1-EXTRACT-ING-026", extract.run),
            Stage("chunk", "M1-INDEX-ING-031", chunk_module.run),
            Stage("embed", "M1-INDEX-ING-032"),
        ),
    )
    yield settings


def _write_docx_with_a_rate_table(path: Path) -> None:
    document = docx.Document()
    document.add_heading("Renewal Terms", level=1)
    document.add_paragraph("Either party may terminate on ninety days written notice.")
    table = document.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "Tier"
    table.rows[0].cells[1].text = "Monthly rate"
    table.rows[1].cells[0].text = "Standard"
    table.rows[1].cells[1].text = "199.00"
    table.rows[2].cells[0].text = "Premium"
    table.rows[2].cells[1].text = "349.00"
    document.save(str(path))


async def _process(
    factory: async_sessionmaker[AsyncSession],
    settings: Settings,
    session: AsyncSession,
    tmp_path: Path,
    filename: str,
) -> tuple[str, list[tuple], uuid.UUID]:
    await nominate(session, str(tmp_path))
    documents = await recorded(session, tmp_path, filename)
    document_id = documents[0]
    outcome = await ingest.process(factory, settings, document_id)
    chunks = (
        await session.execute(
            text(
                "SELECT ordinal, page_from, page_to, heading, content FROM chunks "
                "WHERE document_id = :id ORDER BY ordinal"
            ),
            {"id": document_id},
        )
    ).all()
    return outcome, [tuple(row) for row in chunks], document_id


async def test_a_retrieved_row_carries_its_column_headings(
    factory: async_sessionmaker[AsyncSession],
    session: AsyncSession,
    tmp_path: Path,
    unreachable_queue: Settings,
) -> None:
    """The ticket's own headline scenario: a table row does not arrive
    without the header that gives its number meaning."""
    _write_docx_with_a_rate_table(tmp_path / "rates.docx")
    outcome, chunks, _ = await _process(factory, unreachable_queue, session, tmp_path, "rates.docx")

    assert outcome == "parked"  # extract and chunk succeeded; embed is not installed yet
    assert len(chunks) >= 1

    table_chunk = next(chunk for chunk in chunks if "[TABLE]" in chunk[4])
    assert "Tier | Monthly rate" in table_chunk[4]
    assert "Premium | 349.00" in table_chunk[4]


async def test_ordinals_are_sequential_and_document_order_is_preserved(
    factory: async_sessionmaker[AsyncSession],
    session: AsyncSession,
    tmp_path: Path,
    unreachable_queue: Settings,
) -> None:
    _write_docx_with_a_rate_table(tmp_path / "rates.docx")
    _, chunks, _ = await _process(factory, unreachable_queue, session, tmp_path, "rates.docx")

    assert [chunk[0] for chunk in chunks] == list(range(len(chunks)))


async def test_a_heading_is_recorded_on_the_chunk_beneath_it(
    factory: async_sessionmaker[AsyncSession],
    session: AsyncSession,
    tmp_path: Path,
    unreachable_queue: Settings,
) -> None:
    _write_docx_with_a_rate_table(tmp_path / "rates.docx")
    _, chunks, _ = await _process(factory, unreachable_queue, session, tmp_path, "rates.docx")

    assert any(chunk[3] == "Renewal Terms" for chunk in chunks)


async def test_every_chunk_records_a_page_range_and_no_chunk_is_empty(
    factory: async_sessionmaker[AsyncSession],
    session: AsyncSession,
    tmp_path: Path,
    unreachable_queue: Settings,
) -> None:
    _write_docx_with_a_rate_table(tmp_path / "rates.docx")
    _, chunks, _ = await _process(factory, unreachable_queue, session, tmp_path, "rates.docx")

    for _ordinal, page_from, page_to, _heading, content in chunks:
        assert page_from is not None
        assert page_to is not None
        assert page_from <= page_to
        assert content.strip() != ""


async def test_a_plain_text_document_with_no_headings_still_chunks_by_size(
    factory: async_sessionmaker[AsyncSession],
    session: AsyncSession,
    tmp_path: Path,
    unreachable_queue: Settings,
) -> None:
    sentence = "Either party may terminate this agreement on ninety days written notice. "
    (tmp_path / "note.txt").write_text(sentence * 60)
    outcome, chunks, _ = await _process(factory, unreachable_queue, session, tmp_path, "note.txt")

    assert outcome == "parked"
    assert len(chunks) >= 1
    for chunk in chunks:
        content = chunk[4]
        assert len(content) <= 2400
        assert content.strip() != ""


async def test_re_running_chunk_replaces_rather_than_duplicates(
    factory: async_sessionmaker[AsyncSession],
    session: AsyncSession,
    tmp_path: Path,
    unreachable_queue: Settings,
) -> None:
    """A retried job must not leave two generations of chunks behind it —
    the same idempotency `extract_common.write_anchors` already guarantees
    for `document_pages`."""
    _write_docx_with_a_rate_table(tmp_path / "rates.docx")
    await nominate(session, str(tmp_path))
    documents = await recorded(session, tmp_path, "rates.docx")
    document_id = documents[0]

    work = Work(
        document_id=document_id,
        source_id=uuid.uuid4(),
        path=str(tmp_path / "rates.docx"),
        filename="rates.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        sha256="0" * 64,
    )

    async def report(_done: int, _total: int) -> None:
        return None

    await extract.run(work, report, factory, unreachable_queue)
    await chunk_module.run(work, report, factory, unreachable_queue)
    first_run_count = len(
        (
            await session.execute(
                text("SELECT id FROM chunks WHERE document_id = :id"), {"id": document_id}
            )
        ).all()
    )

    await chunk_module.run(work, report, factory, unreachable_queue)

    async with session_scope(factory) as scoped:
        rows = (
            (
                await scoped.execute(
                    text("SELECT ordinal FROM chunks WHERE document_id = :id ORDER BY ordinal"),
                    {"id": document_id},
                )
            )
            .scalars()
            .all()
        )

    assert len(rows) == first_run_count
    assert rows == list(range(len(rows)))
