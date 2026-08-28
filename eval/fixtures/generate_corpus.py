#!/usr/bin/env python3
"""Builds the grounded-QA fixture corpus committed under `fixtures/corpus/`.

    python eval/fixtures/generate_corpus.py

Deterministic and self-contained — the PDF builders here are trimmed copies
of `api/tests/test_ingest_records.py`'s own `_pdf`/`_scanned_pdf` (kept
separate rather than imported: `eval/` runs outside pytest and must not
depend on `api/tests/` to build its fixtures). Every fact in this corpus is
invented — a fictional company, "Meridian Loom", with fictional policies,
products and figures — so that no task in `grounded_qa.v1.json` can be
answered from a model's general knowledge (`AGENTS.md` §"Validation Rules":
a fixture answerable without the corpus measures the wrong thing).

Run this after editing any fact below; the generated files are committed
so the suite is reproducible without regenerating them, and the fact
strings here must stay byte-for-byte in sync with the substrings named as
`expected_passages` in `eval/suites/grounded_qa.v1.json`.
"""

import io
from pathlib import Path

import docx
import openpyxl
from PIL import Image, ImageDraw, ImageFont

OUT_DIR = Path(__file__).resolve().parent / "corpus"


def _pdf(*pages: str) -> bytes:
    """One page per string, real vector text pdfium can extract — trimmed
    from `api/tests/test_ingest_records.py::_pdf` (no rotation, this corpus
    needs none)."""
    count = len(pages)
    font_number = 3 + count
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        f"<< /Type /Pages /Kids [{' '.join(f'{3 + i} 0 R' for i in range(count))}] "
        f"/Count {count} >>".encode(),
    ]
    for page_index in range(count):
        content_number = font_number + 1 + page_index
        objects.append(
            f"<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 {font_number} 0 R >> >> "
            f"/MediaBox [0 0 612 792] /Contents {content_number} 0 R >>".encode()
        )
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    for page in pages:
        content = f"BT /F1 12 Tf 72 700 Td ({page}) Tj ET".encode("latin-1")
        objects.append(b"<< /Length %d >>\nstream\n" % len(content) + content + b"\nendstream")

    body = bytearray(b"%PDF-1.7\n")
    offsets = []
    for number, obj in enumerate(objects, start=1):
        offsets.append(len(body))
        body += f"{number} 0 obj\n".encode() + obj + b"\nendobj\n"
    xref_offset = len(body)
    body += f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode()
    for offset in offsets:
        body += f"{offset:010d} 00000 n \n".encode()
    body += f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n".encode()
    body += f"startxref\n{xref_offset}\n%%EOF".encode()
    return bytes(body)


def _scanned_pdf(*lines: str) -> bytes:
    """One page, no text layer — an image XObject only, forcing the OCR
    fallback (`extract_ocr.py`). Trimmed from
    `api/tests/test_ingest_records.py::_scanned_pdf` (no rotation)."""
    image = Image.new("RGB", (1400, 300 + 60 * max(len(lines) - 1, 0)), "white")
    draw = ImageDraw.Draw(image)
    font = ImageFont.load_default(size=36)
    for index, line in enumerate(lines):
        draw.text((20, 40 + index * 60), line, fill="black", font=font)

    buffer = io.BytesIO()
    image.save(buffer, format="JPEG", quality=90)
    jpeg = buffer.getvalue()
    width, height = image.size

    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        f"<< /Type /Page /Parent 2 0 R /Resources << /XObject << /Im1 5 0 R >> >> "
        f"/MediaBox [0 0 {width} {height}] /Contents 4 0 R >>".encode(),
    ]
    content = f"q {width} 0 0 {height} 0 0 cm /Im1 Do Q".encode()
    objects.append(b"<< /Length %d >>\nstream\n" % len(content) + content + b"\nendstream")
    objects.append(
        f"<< /Type /XObject /Subtype /Image /Width {width} /Height {height} "
        f"/ColorSpace /DeviceRGB /BitsPerComponent 8 /Filter /DCTDecode "
        f"/Length {len(jpeg)} >>\nstream\n".encode()
        + jpeg
        + b"\nendstream"
    )

    body = bytearray(b"%PDF-1.7\n")
    offsets = []
    for number, obj in enumerate(objects, start=1):
        offsets.append(len(body))
        body += f"{number} 0 obj\n".encode() + obj + b"\nendobj\n"
    xref_offset = len(body)
    body += f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode()
    for offset in offsets:
        body += f"{offset:010d} 00000 n \n".encode()
    body += f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n".encode()
    body += f"startxref\n{xref_offset}\n%%EOF".encode()
    return bytes(body)


# --- content, byte-for-byte matched by eval/suites/grounded_qa.v1.json -----

HANDBOOK_A_PAGES = [
    "Meridian Loom employees accrue eleven paid holiday days per calendar year.",
    "The standard notice period for resignation at Meridian Loom is sixty-three days.",
    "Meridian Loom's parental leave benefit is twenty-two weeks at full pay.",
    "Employees may carry over at most nine unused holiday days into the next year.",
    "The Meridian Loom probation period for new hires lasts one hundred and four days.",
    "Remote employees at Meridian Loom must attend the Cedarbrook office in person "
    "four times per year.",
    "Meridian Loom reimburses home office equipment up to four hundred and forty "
    "dollars per employee.",
    "The Meridian Loom whistleblower hotline number is 800-555-0177.",
]

HANDBOOK_B_PAGES = [
    "Meridian Loom requires laptop disk encryption using a passphrase of at least "
    "sixteen characters.",
    "VPN access tokens at Meridian Loom expire after fourteen days of inactivity.",
    "The Meridian Loom incident response team must be notified within three hours "
    "of a suspected breach.",
    "Meridian Loom rotates production database credentials every forty-five days.",
    "Only the Cedarbrook data center may host Meridian Loom's customer database.",
    "Meridian Loom's password policy requires a minimum length of eighteen "
    "characters with no reuse of the last ten passwords.",
    "Contractors at Meridian Loom receive network access for a maximum of one "
    "hundred and twenty days before renewal.",
    "The Meridian Loom bug bounty program pays a minimum of two thousand dollars "
    "for a critical finding.",
]

NOTICE_SCAN_LINES = [
    "Meridian Loom's annual staff retreat will be held in Alderford this year.",
    "The retreat runs from the ninth to the twelfth of March.",
    "All department leads must confirm attendance by the twentieth of February.",
    "Meridian Loom will cover travel costs up to seven hundred dollars per attendee.",
    "Meridian Loom's whistleblower hotline number is 800-555-0177, the same one "
    "printed in the handbook.",
]

# (heading, fact) pairs — one heading per fact so the docx chunker sees clean
# boundaries between them.
SPEC_SECTIONS = [
    ("Battery Life", "The Loomwear Sensor Mk3 operates for eleven hours on a single charge."),
    ("Weatherproofing", "Its weatherproof rating is IP67."),
    (
        "Connectivity",
        "The sensor connects over Bluetooth 5.3 within a range of forty meters.",
    ),
    ("Firmware Updates", "Firmware updates for the Mk3 are released every six weeks."),
    ("Warranty", "The Mk3's warranty period is three years from the date of purchase."),
    ("Packaging", "Loomwear Sensor Mk3 units ship in packs of six."),
]

# (department, Q1 revenue USD, headcount, avg tenure years)
FIGURES_ROWS = [
    ("Textiles", 482000, 34, 4.1),
    ("Logistics", 215000, 19, 2.7),
    ("Research", 903000, 27, 5.6),
    ("Retail", 118000, 41, 1.9),
    ("Design", 76000, 12, 3.3),
]

# --- conflicting-source pairs, `eval/suites/conflicting_sources.v1.json` ----
#
# `CONFLICT_2025_PAGES`/`CONFLICT_2026_PAGES` are two documents that both
# stay live (neither supersedes the other), dated in their own text, each
# restating the same fact as the other with a genuinely different value —
# real conflicts the answer must present as both positions, never one. The
# "gift card" pair is the deliberate exception: same value, different
# wording, the ticket's own false-conflict-on-wording edge case. The
# `STORE_HOURS` pair is a *separate* two-document case where the eval
# harness itself marks the 2025 page superseded by the 2026 one
# (`eval/conflict.py`'s `_ensure_superseded`) — retrieval then never sees
# both at once, so it is not a conflict task's job to detect anything there.

CONFLICT_2025_PAGES = [
    "As of March 2025, Meridian Loom's standard product return window is "
    "thirty days.",
    "As of March 2025, the express shipping fee at Meridian Loom is twelve "
    "dollars per order.",
    "As of March 2025, Meridian Loom's loyalty program awards one point per "
    "ten dollars spent.",
    "As of March 2025, the maximum trade-in credit for a Loomwear Sensor is "
    "fifty dollars.",
    "As of March 2025, a Meridian Loom gift card balance never expires.",
    "As of March 2025, a Meridian Loom customer support ticket must receive "
    "a first response within twenty-four hours.",
    "As of March 2025, Meridian Loom's affiliate commission rate is eight "
    "percent.",
    "As of March 2025, Meridian Loom warehouses restock inventory every "
    "Tuesday.",
    "As of March 2025, a Meridian Loom extended warranty costs thirty "
    "dollars annually.",
]

CONFLICT_2026_PAGES = [
    "As of January 2026, Meridian Loom's standard product return window is "
    "forty-five days.",
    "As of January 2026, the express shipping fee at Meridian Loom is "
    "eighteen dollars per order.",
    "As of January 2026, Meridian Loom's loyalty program awards one point "
    "per eight dollars spent.",
    "As of January 2026, the maximum trade-in credit for a Loomwear Sensor "
    "is seventy-five dollars.",
    "As of January 2026, gift cards issued by Meridian Loom never expire, "
    "matching the policy stated the previous year.",
    "As of January 2026, a Meridian Loom customer support ticket must "
    "receive a first response within twelve hours.",
    "As of January 2026, Meridian Loom's affiliate commission rate is "
    "eleven percent.",
    "As of January 2026, Meridian Loom warehouses restock inventory every "
    "Thursday.",
    "As of January 2026, a Meridian Loom extended warranty costs "
    "forty-five dollars annually.",
]

STORE_HOURS_2025_LINE = "Meridian Loom retail stores close at 8 PM on weekdays."
STORE_HOURS_2026_LINE = "Meridian Loom retail stores close at 9 PM on weekdays."


def build_handbook_a() -> bytes:
    return _pdf(*HANDBOOK_A_PAGES)


def build_handbook_b() -> bytes:
    return _pdf(*HANDBOOK_B_PAGES)


def build_notice_scan() -> bytes:
    return _scanned_pdf(*NOTICE_SCAN_LINES)


def build_spec_docx() -> bytes:
    document = docx.Document()
    document.add_heading("Product Specification: Loomwear Sensor Mk3", level=0)
    for heading, fact in SPEC_SECTIONS:
        document.add_heading(heading, level=1)
        document.add_paragraph(fact)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_conflict_2025() -> bytes:
    return _pdf(*CONFLICT_2025_PAGES)


def build_conflict_2026() -> bytes:
    return _pdf(*CONFLICT_2026_PAGES)


def build_store_hours_2025() -> bytes:
    return _pdf(STORE_HOURS_2025_LINE)


def build_store_hours_2026() -> bytes:
    return _pdf(STORE_HOURS_2026_LINE)


def build_figures_xlsx() -> bytes:
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Quarterly Department Figures"
    sheet.append(["Department", "Q1 Revenue", "Headcount", "Avg Tenure Years"])
    for row in FIGURES_ROWS:
        sheet.append(list(row))
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    files = {
        "handbook_a.pdf": build_handbook_a(),
        "handbook_b.pdf": build_handbook_b(),
        "notice_scan.pdf": build_notice_scan(),
        "spec.docx": build_spec_docx(),
        "figures.xlsx": build_figures_xlsx(),
        "conflict_2025.pdf": build_conflict_2025(),
        "conflict_2026.pdf": build_conflict_2026(),
        "store_hours_2025.pdf": build_store_hours_2025(),
        "store_hours_2026.pdf": build_store_hours_2026(),
    }
    for name, data in files.items():
        (OUT_DIR / name).write_bytes(data)
        print(f"wrote {OUT_DIR / name} ({len(data)} bytes)")  # noqa: T201 - a build script


if __name__ == "__main__":
    main()
